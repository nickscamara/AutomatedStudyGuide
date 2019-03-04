import docx
from docx import Document
import wikipedia
from docx import Document
from docx.shared import Inches

def addHeader(doc,count):
    if count == 0:
        pp = doc.add_paragraph()
        aa = pp.add_run("Study Guide Generated Below:")
        aa.bold = True
        aa.underline = True
        count += 1
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        try:
            if(para.text != "Terms:"):
                p = doc.add_paragraph()
                runn = p.add_run("-" + para.text)
                runn.bold = True
                doc.add_paragraph(wikipedia.summary(para.text,sentences=1))
        except:
            print("")
        doc.save('static/new.docx')
    return '\n'.join(fullText)

studyGuide = getText('study.docx')

print(studyGuide)