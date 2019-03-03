import wikipedia
import urllib.request
from svglib.svglib import svg2rlg
from reportlab.graphics import renderPDF, renderPM
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from flask import send_file
from watson_developer_cloud import VisualRecognitionV3 as vr
import requests
import json
import pprint
def fontArial():
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Tahoma'
    font.size = Pt(11)  
    
def downloadImage(url,file_path,file_name):
    path = file_path + file_name + url[-4:]
    urllib.request.urlretrieve(url, path)
    if url[-4:] == ".svg":
        drawing = svg2rlg(path)
        renderPM.drawToFile(drawing, 'images/'+file_name+'.png', fmt="PNG")
        return ".png"
    return url[-4:]

def initiliazeWatson():
    ibmWatson = vr(iam_apikey='qqBbMGQ4qmRaPBLbGENUrJMtt-Xy3PvxQk_sptgYDCzJ',version='2016-05-20')
    return ibmWatson

def searchImage(url,facial):
    watson = initiliazeWatson()
    if facial == True:
        imgW = watson.detect_faces(url=url)
        pprint.pprint(imgW.result['images'][0]['faces'])
        return imgW.result['images'][0]['faces']
    else:
        imgW = watson.classify(url=url)
       
        pprint.pprint(imgW.result['images'][0]['classifiers'][0]['classes'][0]['class'])
        return imgW.result['images'][0]['classifiers'][0]['classes'][0]['class']

def searchImageByFileName(url):
    watson = initiliazeWatson()
    imgW = watson.classify(images_file=url)
    pprint.pprint(imgW.result['images'][0]['classifiers'][0]['classes'][0]['class'])
    return imgW.result['images'][0]['classifiers'][0]['classes'][0]['class']

def testing():
    print("it worked")

#Main function
def mainfunction(thename,number,sub1,sub2):
    for x in range(1):

        print(thename)
        #nameOfImage = searchImage("https://pbs.twimg.com/profile_images/988775660163252226/XpgonN0X_400x400.jpg",True)

    
        
        #nameOfImage = searchImage('http://www.unh.edu/unhtales/wp-content/uploads/2014/04/why-i-chose-unh.jpg')
        
       
        #search = input("Enter a word that you would like to create a Study Guide: ")
        search = thename
        num = number
        #num = input("Enter the number of sub-topics that you would like to have: ")
        array = []
        nums = int(num)
        '''
        iterator = 0
        while iterator < nums:
            addthis = input("Enter a sub-topic: ")
            array.append(addthis)
            iterator = iterator + 1
        print(array)
        '''

        array.append(sub1)
        array.append(sub2)


        doc = Document()
        #fontArial()
        doc.add_heading(search.capitalize(),0)
        wiki = search
        last_paragraph = doc.add_paragraph((wikipedia.summary(wiki,sentences=3)))
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph.style = doc.styles['Normal']
        
        for y in array:
            count = 1
            doc.add_heading(y.capitalize(), level=1)
            last_paragraph = doc.add_paragraph((wikipedia.summary(y)))
            
            #wikipedia.random()
            save_in_array = []
            getUrls = []

            
            br = wikipedia.page(y)
            ite = 0
        

            for imag in br.images:
                url = br.images[ite]
                if url[-4:] != '.svg':
                    print("here")
                    print(br.images[ite])
                    file_name = br.title.strip()
                    aa = file_name.replace(" ", "")
                    save_in_array.append(aa)
                    #print(aa)
                    urls = downloadImage(url,'images/', save_in_array[ite]+str(ite))
                    getUrls.append(urls)
                    print(urls)
                    ite += 1
                
            ite = 0
            for img in save_in_array:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                picture = run.add_picture('images/'+ save_in_array[ite]+str(ite) + getUrls[ite], width=Inches(3))
                print(save_in_array[ite] + urls)
                paragraph = doc.paragraphs[-1] 
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ite += 1
            
            
            '''if url[-4:] != ".svg":
                nameImage = searchImage(url)
                print("The AI Name of the Image: " + nameImage)
                subtitle = doc.add_paragraph(str(count) + ". " + nameImage )
                subtitle = doc.paragraphs[-1] 
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                urlimg = 'images/'+ aa + urls
                print(urlimg + " this is the img url")
                nameImage = searchImageByFileName(urlimg)
                print("The AI Name of the Image: " + nameImage)
                subtitle = doc.add_paragraph(str(count) + ". " + nameImage )
                subtitle = doc.paragraphs[-1] 
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER'''

            #img = doc.add_picture('images/'+ aa + urls, width=Inches(3))
            #img.alignment  = 1
            count += 1
        
        doc.save('static/me.docx')


